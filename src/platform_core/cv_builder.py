"""
Builds a professional CV (.docx) from a user's Profile data: personal info,
employment history, education, skills, and references. Used by the
POST /cv/generate route in main.py.
"""
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ACCENT_COLOR = RGBColor(0x1D, 0x4E, 0xD8)  # blue-700, matches the platform's brand color


def _heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT_COLOR
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    # Simple bottom border under the heading for visual separation
    p_pr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1D4ED8')
    borders.append(bottom)
    p_pr.append(borders)
    return p


def build_cv_docx(user, employment_history, education, references, output_path: str):
    doc = DocxDocument()

    # Base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # --- Header: name + contact line ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(user.full_name or "Full Name Not Set")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = ACCENT_COLOR

    contact_bits = [b for b in [user.email, user.phone, user.address, user.linkedin_url] if b]
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_p.add_run(" | ".join(contact_bits))
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Professional identity / right to work ---
    identity_bits = []
    if user.professional_registration_body and user.professional_registration_number:
        identity_bits.append(f"{user.professional_registration_body} Registration: {user.professional_registration_number}")
    if user.visa_status:
        identity_bits.append(f"Right to Work: {user.visa_status}")
    if identity_bits:
        identity_p = doc.add_paragraph()
        identity_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        identity_run = identity_p.add_run(" | ".join(identity_bits))
        identity_run.italic = True
        identity_run.font.size = Pt(9.5)

    # --- Professional Summary ---
    if user.experience_summary:
        _heading(doc, "Professional Summary")
        doc.add_paragraph(user.experience_summary)

    # --- Employment History ---
    if employment_history:
        _heading(doc, "Employment History")
        for job in employment_history:
            p = doc.add_paragraph()
            title_run = p.add_run(job.job_title)
            title_run.bold = True
            if job.grade_band:
                p.add_run(f"  ·  {job.grade_band}").italic = True

            dates = f"{job.start_date or ''} – {'Present' if job.is_current else (job.end_date or '')}"
            date_p = doc.add_paragraph()
            date_p.paragraph_format.space_after = Pt(2)
            sub_run = date_p.add_run(f"{job.institution_name}" + (f", {job.location}" if job.location else "") + f"   ({dates})")
            sub_run.font.size = Pt(9.5)
            sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Education & Qualifications ---
    if education:
        _heading(doc, "Education & Qualifications")
        for edu in education:
            p = doc.add_paragraph()
            p.add_run(edu.qualification_name).bold = True
            details = " · ".join([b for b in [edu.institution, edu.qualification_type, edu.date_awarded] if b])
            if details:
                p.add_run(f"  —  {details}").font.size = Pt(9.5)

    # --- Clinical Skills & Competencies ---
    if user.skills_list:
        _heading(doc, "Clinical Skills & Competencies")
        skills = [s.strip() for s in user.skills_list.split(",") if s.strip()]
        doc.add_paragraph(" • " + "   •  ".join(skills))

    # --- Declarations ---
    declarations = []
    if user.dbs_status:
        declarations.append(f"DBS: {user.dbs_status}")
    if user.health_declaration_status:
        declarations.append(f"Health Declaration: {user.health_declaration_status}")
    if user.indemnity_status:
        declarations.append(f"Professional Indemnity: {user.indemnity_status}")
    if declarations:
        _heading(doc, "Declarations")
        doc.add_paragraph(" | ".join(declarations))

    # --- References ---
    if references:
        _heading(doc, "References")
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        for i, label in enumerate(["Name", "Role / Institution", "Email", "Phone"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].bold = True
        for ref in references:
            row = table.add_row().cells
            row[0].text = ref.name or ""
            row[1].text = " / ".join([b for b in [ref.role, ref.institution] if b])
            row[2].text = ref.email or ""
            row[3].text = ref.phone or ""
    else:
        _heading(doc, "References")
        doc.add_paragraph("Available on request.")

    doc.save(output_path)
    return output_path
